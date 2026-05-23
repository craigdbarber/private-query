"""Contains utility functions for text transformation."""

import os
from collections.abc import Callable
from pathlib import Path

import docx
from loguru import logger
from pypdf import PdfReader


def extract_text_from_file(file: str | Path) -> str:
    """Read the specified file and extract all readable text.

    Args:
        file: The file to be read.

    Returns: The extracted text.

    Raises:
        RuntimeError: If the file type is not supported.
        ValueError: If the file path is not a file, or the file type
        is not supported.
        FileNotFoundError: If the file does not exist.
        PermissionError: If the file could not be opened.

    """
    file_extractors: dict[str, Callable[[str | Path], str]] = {
        ".pdf": extract_text_from_pdf,
        ".docx": extract_text_from_docx,
        ".txt": extract_text_from_txt,
    }
    file_path = Path(file)

    # error checking
    if not file_path.exists():
        err_msg = f"File does not exist: {file}"
        logger.error(err_msg)
        raise FileNotFoundError(err_msg)
    if not file_path.is_file():
        err_msg = f"File path is not file: {file}"
        logger.error(err_msg)
        raise ValueError(err_msg)
    if not os.access(file_path, mode=os.R_OK):
        err_msg = f"Could not open file: {file}"
        logger.error(err_msg)
        raise PermissionError(err_msg)
    if file_path.suffix not in file_extractors:
        err_msg = f"Unsupported file type: {file_path.suffix}"
        logger.error(err_msg)
        raise ValueError(err_msg)

    logger.info(f"Extracting text from file: {file}")
    return file_extractors[file_path.suffix](file_path)


def extract_text_from_txt(txt_path: str | Path) -> str:
    """Read the specified file and extract all text.

    Args:
        txt_path: The path of the txt file to be read.

    Returns: The extracted text.

    """
    with open(txt_path, encoding="UTF-8") as file:
        return file.read()


def extract_text_from_docx(docx_path: str | Path) -> str:
    """Read the specified docx file and extract all readable text.

    Args:
        docx_path: The path of the docx file to be read.
    Returns: The extracted readable text.

    """
    file = Path(docx_path)
    if not file.exists():
        raise FileNotFoundError(f"The docx path: {docx_path} does not exists.")

    doc = docx.Document(str(file))
    extracted_text: list[str] = []
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            extracted_text.append(paragraph.text)

    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_text:
                extracted_text.append("\t".join(row_text))

    return "\n\n".join(extracted_text)


def extract_text_from_pdf(pdf_path: str | Path) -> str:
    """Read the specified pdf file and extract all readable text.

    Args:
        pdf_path: The path of the pdf file to be read.
    Returns: The extracted readable text.

    Raises:
        FileNotFoundError: If the specified pdf path doesn't exist.

    """
    file = Path(pdf_path)
    if not file.exists():
        raise FileNotFoundError(f"The pdf path: {pdf_path} does not exist.")
    reader = PdfReader(file)
    text = ""
    for page in reader.pages:
        text += page.extract_text() or ""
    return text
